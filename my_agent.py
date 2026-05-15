import anthropic
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

client = anthropic.Anthropic()

# ============================================
# KELVIN'S REAL BACKGROUND
# ============================================
KELVIN_CV = """
Name: Kelvin Wong, PhD
Contact: Pittsburg, CA | 510-605-8602 | kelvinmsu@outlook.com

EXPERIENCE:
- Genomics Epidemiologist, California Dept of Public Health (CDPH), 2022-present
  * Wastewater surveillance across 80+ California sites
  * Built R Shiny dashboard for COVID variant tracking
  * Applied generative AI tools to automate workflows
  * Automated data processes using DevOps and RStudio Connect

- City Research Scientist, NYC Dept of Health, 2020-2022
  * Managed lab processing 5000+ vector samples/year
  * Developed multiplex panel qPCR assay for vector-borne pathogens
  * Validated molecular assays

- Senior Scientist, Sequlite Genomics, 2019-2020
  * Optimized NGS reagent chemistry
  * Built R code pipeline for sequencing data analysis
  * Developed clustering algorithm improving detection accuracy by 30%

- Post-Doctoral Scientist, EPA, 2010-2016
  * EPA Level 1 Science and Technology Award
  * Metagenomics sequencing, qPCR, bioinformatics pipeline in bash/R/Python

- Research Assistant, Michigan State University, 2006-2010
  * Developed multiplex qPCR assays for fecal source tracking
  * Pioneered wastewater viral testing methods used in COVID monitoring

EDUCATION:
- PhD Environmental Engineering, Michigan State University, 2010
- MS Environmental Engineering, Southern Illinois University, 2000
- BA Biochemistry, Southern Illinois University, 1998

SKILLS:
- Python (Pandas, NumPy, Scikit-Learn, SciPy, Matplotlib)
- R (dplyr, ggplot2, Shiny, rmarkdown, bioconductor)
- Machine learning: Random Forest, clustering, PCA, time series, anomaly detection
- NGS, metagenomics, qPCR, library prep, bioinformatics pipelines
- SQL, bash, Git, DevOps, RStudio Connect
- Currently building AI agents with Claude API and Claude Code

AWARDS: EPA Level 1 Science Award 2017, EPA Honorable Mention 2014
PUBLICATIONS: 20+ peer-reviewed papers, ~1000 citations
"""

# ============================================
# PASTE YOUR REAL JOB HERE
# ============================================
JOB_TITLE = "Applied AI Scientist - Verily (Google Health)"
JOB_DESCRIPTION = """
We are looking for an Applied AI Scientist to develop ML models 
for population health applications. You will work on biosurveillance 
systems, build production ML pipelines, and translate complex health 
data into actionable insights.

Requirements:
- PhD in relevant field preferred
- Python, machine learning, production ML pipelines
- Experience with genomics or bioinformatics a plus
- Public health or clinical data experience
- Strong publication record
"""

print("=" * 60)
print("🤖 KELVIN'S JOB SEARCH AGENT")
print("=" * 60)

# ============================================
# STEP 1 — Score the job (Lesson 3)
# ============================================
print(f"\n📊 STEP 1: Scoring job fit for {JOB_TITLE}...")

response = client.messages.create(
    model="claude-sonnet-4-5",
    max_tokens=1024,
    system="You are a job fit analyzer. Respond ONLY with valid JSON. No markdown.",
    messages=[{
        "role": "user",
        "content": f"""
Analyze this job for Kelvin:
BACKGROUND: {KELVIN_CV}
JOB: {JOB_TITLE}
REQUIREMENTS: {JOB_DESCRIPTION}

Return JSON with:
- fit_score (0-100)
- apply_recommendation (yes/no/maybe)
- top_3_strengths (list)
- gaps (list, empty if none)
- one_line_reason (string)
"""
    }]
)

raw = response.content[0].text.replace("```json","").replace("```","").strip()
score = json.loads(raw)

print(f"  Fit Score:  {score['fit_score']}/100")
print(f"  Apply?      {score['apply_recommendation'].upper()}")
print(f"  Why:        {score['one_line_reason']}")
print(f"  Strengths:  {score['top_3_strengths']}")
print(f"  Gaps:       {score['gaps'] or 'None!'}")

if score['apply_recommendation'] == 'no':
    print("\n⚠️  Low fit score — consider finding a better matched job.")
    exit()

# ============================================
# STEP 2 — Multi-turn: tailor resume + cover letter
# (Lesson 4)
# ============================================
print(f"\n✍️  STEP 2: Tailoring resume and cover letter...")

conversation = []

def chat(message):
    conversation.append({"role": "user", "content": message})
    response = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=2000,
        system=f"""You are an expert resume writer helping Kelvin apply for: 
{JOB_TITLE}. His background: {KELVIN_CV}. 
Always be specific, use metrics, and tailor everything to the job.""",
        messages=conversation
    )
    reply = response.content[0].text
    conversation.append({"role": "assistant", "content": reply})
    return reply

# Turn 1 — get tailored resume content as JSON
print("  → Generating tailored resume content...")
turn1 = chat(f"""
For this job: {JOB_TITLE}
Requirements: {JOB_DESCRIPTION}

Give me tailored resume content. Respond ONLY in valid JSON with:
- summary (2 punchy sentences)
- skills (list of 8 most relevant)
- experience (list of 4, each with: title, company, dates, bullets (2 strong bullets with metrics))
- education (list of 3, each with: degree, school, year) 
- awards (list of 2)
""")

resume_data = json.loads(turn1.replace("```json","").replace("```","").strip())

# Turn 2 — get cover letter (remembers job + resume from Turn 1)
print("  → Writing cover letter...")
cover_letter = chat(f"""
Using the resume content you just created, write a professional 
3-paragraph cover letter for the {JOB_TITLE} role.
- Paragraph 1: Hook + why this role
- Paragraph 2: Top 2 specific achievements from my background  
- Paragraph 3: Forward-looking close
Keep it under 250 words. No placeholders.
""")

# ============================================
# STEP 3 — Build Word document (Lesson 6)
# ============================================
print(f"\n📄 STEP 3: Building Word document...")

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin = Pt(36)
section.bottom_margin = Pt(36)
section.left_margin = Pt(54)
section.right_margin = Pt(54)

# --- Name ---
name_para = doc.add_paragraph()
name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = name_para.add_run("KELVIN WONG, PhD")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

# --- Contact ---
contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
c = contact.add_run("Pittsburg, CA  |  510-605-8602  |  kelvinmsu@outlook.com  |  linkedin.com/in/kelvin-wong-13775115")
c.font.size = Pt(9)
c.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# Divider
div = doc.add_paragraph()
div_run = div.add_run("─" * 80)
div_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    # Add underline effect via border
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# --- Summary ---
add_section_heading(doc, "PROFESSIONAL SUMMARY")
p = doc.add_paragraph()
p.add_run(resume_data["summary"]).font.size = Pt(10)

# --- Skills ---
add_section_heading(doc, "TECHNICAL SKILLS")
p = doc.add_paragraph()
p.add_run(" • ".join(resume_data["skills"])).font.size = Pt(10)

# --- Experience ---
add_section_heading(doc, "PROFESSIONAL EXPERIENCE")
for job in resume_data["experience"]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r1 = p.add_run(f"{job['title']}  —  {job['company']}")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    r2 = p.add_run(f"   |   {job['dates']}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    for bullet in job["bullets"]:
        b = doc.add_paragraph(style="List Bullet")
        b.paragraph_format.left_indent = Pt(18)
        br = b.add_run(bullet)
        br.font.size = Pt(10)

# --- Education ---
add_section_heading(doc, "EDUCATION")
for edu in resume_data["education"]:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{edu['degree']}")
    r1.bold = True
    r1.font.size = Pt(10)
    p.add_run(f"  —  {edu['school']},  {edu['year']}").font.size = Pt(10)

# --- Awards ---
add_section_heading(doc, "HONORS & AWARDS")
for award in resume_data["awards"]:
    b = doc.add_paragraph(style="List Bullet")
    b.add_run(award).font.size = Pt(10)

# Save resume
resume_path = "/Users/chiyuenwong/resume_tailored.docx"
doc.save(resume_path)
print(f"  ✅ Resume saved: {resume_path}")

# ============================================
# STEP 4 — Save cover letter as Word doc
# ============================================
cover_doc = Document()
section = cover_doc.sections[0]
section.top_margin = Pt(72)
section.bottom_margin = Pt(72)
section.left_margin = Pt(72)
section.right_margin = Pt(72)

# Header
p = cover_doc.add_paragraph()
p.add_run("Kelvin Wong, PhD").bold = True
cover_doc.add_paragraph("Pittsburg, CA  |  510-605-8602  |  kelvinmsu@outlook.com")
cover_doc.add_paragraph("")

# Job title
p = cover_doc.add_paragraph()
p.add_run(f"Re: {JOB_TITLE}").bold = True
cover_doc.add_paragraph("")

# Cover letter body — split into paragraphs
for para in cover_letter.strip().split("\n\n"):
    if para.strip():
        cover_doc.add_paragraph(para.strip())

# Sign off
cover_doc.add_paragraph("")
cover_doc.add_paragraph("Sincerely,")
cover_doc.add_paragraph("")
p = cover_doc.add_paragraph()
p.add_run("Kelvin Wong, PhD").bold = True

cover_path = "/Users/chiyuenwong/cover_letter.docx"
cover_doc.save(cover_path)
print(f"  ✅ Cover letter saved: {cover_path}")

# ============================================
# DONE!
# ============================================
print("\n" + "=" * 60)
print("🏆 JOB SEARCH AGENT COMPLETE!")
print("=" * 60)
print(f"\n  Job:        {JOB_TITLE}")
print(f"  Fit Score:  {score['fit_score']}/100")
print(f"  Apply?      {score['apply_recommendation'].upper()}")
print(f"\n  📄 Resume:       {resume_path}")
print(f"  ✉️  Cover Letter: {cover_path}")
print("\n  Open them:")
print("  open ~/resume_tailored.docx")
print("  open ~/cover_letter.docx")
print("=" * 60)